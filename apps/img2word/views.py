from __future__ import annotations

import base64
import io
import json
import os

from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.shortcuts import render
from django.views.decorators.http import require_http_methods

PROMPT = """Analise esta imagem cuidadosamente e extraia todo o conteúdo textual preservando a estrutura e formatação.

Retorne SOMENTE um JSON válido (sem markdown, sem explicações) com esta estrutura:
{
  "titulo_documento": "título identificado ou null",
  "elementos": [
    {"tipo": "titulo1", "texto": "Título principal"},
    {"tipo": "titulo2", "texto": "Subtítulo"},
    {"tipo": "titulo3", "texto": "Título nível 3"},
    {"tipo": "paragrafo", "texto": "Texto do parágrafo", "negrito": false, "italico": false},
    {"tipo": "lista_simples", "itens": ["item 1", "item 2"]},
    {"tipo": "lista_numerada", "itens": ["item 1", "item 2"]},
    {"tipo": "tabela", "cabecalhos": ["Col A", "Col B"], "linhas": [["valor1", "valor2"]]},
    {"tipo": "separador"},
    {"tipo": "assinatura", "texto": "Nome / cargo"}
  ]
}

Regras:
- Preserve a hierarquia dos títulos (titulo1 para o maior, titulo2 para o seguinte)
- Para parágrafos com texto em negrito use "negrito": true; itálico use "italico": true
- Tabelas devem manter a estrutura de linhas e colunas exatamente
- Se a imagem não contiver texto, retorne {"titulo_documento": null, "elementos": []}
- Retorne APENAS o JSON, sem nenhum texto adicional antes ou depois
"""


def _build_docx(structure: dict) -> io.BytesIO:
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    for el in structure.get('elementos', []):
        tipo = el.get('tipo', '')

        if tipo == 'titulo1':
            doc.add_heading(el.get('texto', ''), level=1)

        elif tipo == 'titulo2':
            doc.add_heading(el.get('texto', ''), level=2)

        elif tipo == 'titulo3':
            doc.add_heading(el.get('texto', ''), level=3)

        elif tipo == 'paragrafo':
            p = doc.add_paragraph()
            run = p.add_run(el.get('texto', ''))
            if el.get('negrito'):
                run.bold = True
            if el.get('italico'):
                run.italic = True

        elif tipo == 'lista_simples':
            for item in el.get('itens', []):
                doc.add_paragraph(str(item), style='List Bullet')

        elif tipo == 'lista_numerada':
            for item in el.get('itens', []):
                doc.add_paragraph(str(item), style='List Number')

        elif tipo == 'tabela':
            headers = el.get('cabecalhos', [])
            rows = el.get('linhas', [])
            num_cols = max(
                len(headers),
                max((len(r) for r in rows), default=0),
            )
            if num_cols > 0:
                table = doc.add_table(rows=1 + len(rows), cols=num_cols)
                table.style = 'Table Grid'
                hdr_row = table.rows[0]
                for i, h in enumerate(headers[:num_cols]):
                    cell = hdr_row.cells[i]
                    cell.text = str(h)
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].bold = True
                for r_idx, row_data in enumerate(rows):
                    tbl_row = table.rows[r_idx + 1]
                    for c_idx, val in enumerate(row_data[:num_cols]):
                        tbl_row.cells[c_idx].text = str(val)
                doc.add_paragraph()

        elif tipo == 'separador':
            doc.add_paragraph('─' * 50)

        elif tipo == 'assinatura':
            p = doc.add_paragraph(el.get('texto', ''))
            p.runs[0].italic = True if p.runs else None

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _extrair_json(raw: str) -> dict:
    raw = raw.strip()
    if raw.startswith('```'):
        parts = raw.split('```')
        raw = parts[1] if len(parts) > 1 else raw
        if raw.startswith('json'):
            raw = raw[4:]
    return json.loads(raw.strip())


@login_required
@require_http_methods(['GET', 'POST'])
def index(request):
    if request.method == 'GET':
        return render(request, 'img2word/index.html')

    imagem = request.FILES.get('imagem')
    if not imagem:
        return JsonResponse({'erro': 'Nenhuma imagem enviada.'}, status=400)

    tipos_aceitos = {
        'image/jpeg', 'image/png', 'image/gif',
        'image/webp', 'image/bmp', 'image/tiff',
    }
    if imagem.content_type not in tipos_aceitos:
        return JsonResponse({'erro': 'Tipo de arquivo não suportado.'}, status=400)

    try:
        from openai import OpenAI
        client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

        b64 = base64.b64encode(imagem.read()).decode('utf-8')
        data_url = f'data:{imagem.content_type};base64,{b64}'

        resp = client.chat.completions.create(
            model='gpt-4o',
            messages=[{
                'role': 'user',
                'content': [
                    {'type': 'text', 'text': PROMPT},
                    {'type': 'image_url', 'image_url': {'url': data_url}},
                ],
            }],
            max_tokens=4096,
        )

        raw_content = resp.choices[0].message.content
        structure = _extrair_json(raw_content)

    except json.JSONDecodeError:
        return JsonResponse({'erro': 'A IA retornou um formato inesperado. Tente novamente.'}, status=500)
    except Exception as exc:
        return JsonResponse({'erro': f'Erro ao processar imagem: {exc}'}, status=500)

    try:
        buf = _build_docx(structure)
    except Exception as exc:
        return JsonResponse({'erro': f'Erro ao gerar documento Word: {exc}'}, status=500)

    nome_base = imagem.name.rsplit('.', 1)[0]
    filename = f'{nome_base}.docx'

    response = HttpResponse(
        buf.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response
